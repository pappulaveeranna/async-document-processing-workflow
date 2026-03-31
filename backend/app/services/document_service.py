import os
import re
from datetime import datetime
from sqlalchemy.orm import Session
from app.models.models import Document, ProcessingJob, JobStatus
from app.core.config import settings


def save_document_and_job(db: Session, filename: str, original_filename: str,
                           file_type: str, file_size: int, file_path: str):
    doc = Document(
        filename=filename,
        original_filename=original_filename,
        file_type=file_type,
        file_size=file_size,
        file_path=file_path,
    )
    db.add(doc)
    db.flush()

    job = ProcessingJob(document_id=doc.id)
    db.add(job)
    db.commit()
    db.refresh(doc)
    db.refresh(job)
    return doc, job


def get_jobs(db: Session, status: str = None, search: str = None,
             sort_by: str = "created_at", sort_order: str = "desc",
             skip: int = 0, limit: int = 20):
    query = db.query(ProcessingJob).join(ProcessingJob.document)

    if status:
        query = query.filter(ProcessingJob.status == status)
    if search:
        query = query.filter(Document.original_filename.ilike(f"%{search}%"))

    sort_col = getattr(ProcessingJob, sort_by, ProcessingJob.created_at)
    if sort_order == "desc":
        query = query.order_by(sort_col.desc())
    else:
        query = query.order_by(sort_col.asc())

    total = query.count()
    items = query.offset(skip).limit(limit).all()
    return items, total


def get_job_by_id(db: Session, job_id: str):
    return db.query(ProcessingJob).filter(ProcessingJob.id == job_id).first()


def update_job_status(db: Session, job_id: str, status: JobStatus,
                      stage: str = None, progress: int = None,
                      error: str = None, celery_task_id: str = None):
    job = db.query(ProcessingJob).filter(ProcessingJob.id == job_id).first()
    if not job:
        return None
    job.status = status
    if stage is not None:
        job.current_stage = stage
    if progress is not None:
        job.progress = progress
    if error is not None:
        job.error_message = error
    if celery_task_id is not None:
        job.celery_task_id = celery_task_id
    job.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(job)
    return job


def update_extracted_data(db: Session, job_id: str, data: dict):
    job = db.query(ProcessingJob).filter(ProcessingJob.id == job_id).first()
    if not job:
        return None
    job.extracted_data = data
    job.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(job)
    return job


def finalize_job(db: Session, job_id: str, finalized_data: dict):
    job = db.query(ProcessingJob).filter(ProcessingJob.id == job_id).first()
    if not job:
        return None
    job.finalized_data = finalized_data
    job.status = JobStatus.finalized
    job.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(job)
    return job


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text content from uploaded file."""
    try:
        if file_type == "application/pdf":
            import PyPDF2
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = [page.extract_text() or "" for page in reader.pages]
                text = " ".join(pages_text).strip()
                # Fallback metadata if PDF is image-based / unextractable
                if not text:
                    text = f"PDF document with {len(reader.pages)} page(s). Text extraction not available (image-based or encrypted PDF)."
                return text
        elif file_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            return " ".join(p.text for p in doc.paragraphs if p.text.strip())
        elif file_type.startswith("text/"):
            with open(file_path, "r", errors="ignore") as f:
                return f.read()
    except Exception:
        pass
    return ""


def derive_structured_fields(filename: str, file_type: str, file_size: int, text: str) -> dict:
    """Derive structured fields from document content."""
    words = re.findall(r'\b[a-zA-Z]{4,}\b', text.lower())
    freq = {}
    for w in words:
        freq[w] = freq.get(w, 0) + 1
    stopwords = {"this", "that", "with", "from", "have", "will", "been", "they",
                 "their", "there", "were", "what", "when", "which", "your", "also"}
    keywords = [w for w, _ in sorted(freq.items(), key=lambda x: -x[1])
                if w not in stopwords][:10]

    ext = os.path.splitext(filename)[1].lower()
    category_map = {".pdf": "PDF Document", ".docx": "Word Document",
                    ".txt": "Text File", ".csv": "Spreadsheet", ".md": "Markdown"}
    category = category_map.get(ext, "Unknown")

    summary = text[:1000].strip().replace("\n", " ") if text and len(text.strip()) > 10 else "No extractable text content found in this document."

    return {
        "title": os.path.splitext(filename)[0].replace("_", " ").replace("-", " ").title(),
        "category": category,
        "summary": summary,
        "keywords": keywords,
        "file_metadata": {
            "filename": filename,
            "file_type": file_type,
            "file_size_bytes": file_size,
            "word_count": len(text.split()) if text else 0,
            "char_count": len(text),
        },
        "processed_at": datetime.utcnow().isoformat(),
    }
